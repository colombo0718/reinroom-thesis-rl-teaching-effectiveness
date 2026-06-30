#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_thesis.py — 元智大學碩士論文組裝腳本（格式A）

輸入：md/ 資料夾的章節 md 檔 + images/ 資料夾的圖片
輸出：論文_組裝.docx

【格式對齊原則】md 不管正式格式，腳本負責對齊學校規範。
  - 章節編號：md 的「X.X」「X.X.X」由腳本轉成「第N節」「壹、」
  - 圖表編號：md 的「圖 3-1」依出現順序連續編號為「圖 1」、「圖 2」...
  - 文內圖片參照：「見圖 3-1」也會同步轉換
  - 圖片語法：支援 markdown 的 ![caption](path) 語法

圖片命名規則：images/fig{章}-{圖號}.png
  例：圖 3-1 → images/fig3-1.png
  支援字尾：圖 3-11a → images/fig3-11a.png
"""

import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import latex2mathml.converter

# ── 論文 metadata（封面 / 書名頁 / 審定書填入） ────────────────────────────
THESIS_META = {
    'title_zh': 'Rein Room 強化教室：視覺化互動強化學習教學平台之設計與教學成效評估',
    'title_en': 'Design and Evaluation of Rein Room: A Visual Interactive Platform for Reinforcement Learning Education',
    'student_zh': '趙士豪',
    'student_en': 'Shih-Hao Chao',
    'student_id': 'S1136103',
    'advisor_zh': '黃怡錚',
    'advisor_en': 'Yi-Jheng Huang',                     # 已確認（系所網頁）
    'dept_zh': '資訊工程學系',
    'dept_en': 'Computer Science and Engineering',
    'college_en': 'College of Informatics',             # 已確認（系所網頁職稱為「資訊學院英語學士班主任」）
    'degree_zh': '碩士',
    'degree_en': 'Master of Science',
    'year_roc': '一一五',                                # 民國年
    'month_zh': '六',
    'year_ad': '2026',
    'month_en': 'June',
    'location_en': 'Chungli, Taiwan, Republic of China',
}

# ── LaTeX → OMML 轉換（使用 Word 自帶 MML2OMML.XSL） ──────────────────────
MML2OMML_XSL = Path(r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL')
_xslt = etree.XSLT(etree.parse(str(MML2OMML_XSL))) if MML2OMML_XSL.exists() else None
OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def latex_to_omml(tex, display=False):
    """LaTeX 字串 → OMML <m:oMath> Element；轉換失敗時回傳 None。"""
    if _xslt is None:
        return None
    try:
        mathml = latex2mathml.converter.convert(tex, display='block' if display else 'inline')
        omml_tree = _xslt(etree.fromstring(mathml.encode('utf-8')))
        # XSLT 輸出根節點為 m:oMath
        omml_xml = etree.tostring(omml_tree, xml_declaration=False)
        return etree.fromstring(omml_xml)
    except Exception as e:
        print(f'  ⚠  LaTeX 轉換失敗: {tex[:50]}... ({e})')
        return None

BASE    = Path(__file__).parent

# ── 目錄頁碼對照（由 build_with_toc.py 在 2-pass 之間寫入） ──────────────
TOC_PAGES_FILE = BASE / 'toc_pages.json'
TOC_PAGES = {}
if TOC_PAGES_FILE.exists():
    try:
        with open(TOC_PAGES_FILE, encoding='utf-8') as _f:
            TOC_PAGES = json.load(_f)
    except Exception:
        TOC_PAGES = {}

def _toc_page(key):
    """查 toc_pages.json；找不到回傳預設 '—'。"""
    return str(TOC_PAGES.get(key, '—'))
MD_DIR  = BASE / "md"
IMG_DIR = BASE / "images"
OUTPUT  = BASE / "論文_組裝.docx"

FONT_CN = '標楷體'
FONT_EN = 'Times New Roman'

# ── 格式轉換開關（對齊元智格式A） ─────────────────────────────────────────
CONVERT_SECTION_NUMBERING    = True   # 5.1 → 第一節
CONVERT_SUBSECTION_NUMBERING = True   # 5.1.1 → 壹、
CONVERT_FIGURE_NUMBERING     = True   # 圖 3-1 → 圖 1（全文連續編號）
CONVERT_TABLE_NUMBERING      = True   # 表 3-1 → 表 1（全文連續編號）
NORMALIZE_CAPTION_PUNCT      = True   # 圖 1：xxx / 圖 1 xxx → 圖 1　xxx（全形空格）
ADD_PAGE_NUMBER              = True   # footer 中央阿拉伯數字頁碼
BODY_FIRST_LINE_INDENT_CHARS = 2      # 正文首行縮排字元數

# ── 中文數字對照表 ───────────────────────────────────────────────────────
NUM_CN_SIMPLE = ['', '一','二','三','四','五','六','七','八','九','十',
                 '十一','十二','十三','十四','十五','十六','十七','十八','十九','二十']
NUM_CN_FORMAL = ['', '壹','貳','參','肆','伍','陸','柒','捌','玖','拾',
                 '拾壹','拾貳','拾參','拾肆','拾伍','拾陸','拾柒','拾捌','拾玖','貳拾']

def num_to_cn(n, formal=False):
    table = NUM_CN_FORMAL if formal else NUM_CN_SIMPLE
    return table[n] if 0 < n < len(table) else str(n)

def convert_section_text(text):
    """5.1 RL 概念理解測驗分析 → 第一節　RL 概念理解測驗分析"""
    m = re.match(r'^\d+\.(\d+)\s+(.*)$', text)
    if not m:
        return text
    return f'第{num_to_cn(int(m.group(1)))}節　{m.group(2).strip()}'

def convert_subsection_text(text):
    """4.2.1 實驗設計型態 → 壹、實驗設計型態"""
    m = re.match(r'^\d+\.\d+\.(\d+)\s+(.*)$', text)
    if not m:
        return text
    return f'{num_to_cn(int(m.group(1)), formal=True)}、{m.group(2).strip()}'

# ── 章節順序（可自行調整章名與檔案清單） ────────────────────────────────
CHAPTERS = [
    ("第一章　緒論", [
        "第一章 緒論.md",
    ]),
    ("第二章　文獻探討", [
        "2.1 強化學習基本概念與教育潛力.md",
        "2.2 中小學生 RL 教學案例分析.md",
        "2.3 Gymnasium 平台與標準化環境簡介.md",
        "2.4 視覺化介面（UI-driven）與低門檻 RL 教學.md",
        "2.5 本研究切入點與創新定位.md",
    ]),
    ("第三章　Rein Room 平台設計", [
        "3.1 RL Lab 平台架構.md",
        "3.2 演算法設計.md",
        "3.3 視覺化模組.md",
        "3.4 遊戲環境.md",
    ]),
    ("第四章　研究設計", [
        "4.1 研究假設與實驗目標.md",
        "4.2 A_B 教學實驗流程設計.md",
        "4.3 實驗對象與分組方式說明.md",
        "4.4 教學流程與教材使用說明.md",
        "4.5 評估工具設計.md",
        "4.6 數據分析方法.md",
    ]),
    ("第五章　研究結果與分析", [
        "5.1 RL 概念理解測驗分析_.md",
        "5.2 任務完成成效與策略發展比較.md",
        "5.3 學習參與度與平台使用回饋.md",
        "5.4 錯誤行為觀察與平台改進討論.md",
    ]),
    ("第六章　結論與建議", [
        "6.1 研究成果總結.md",
        "6.2 平台與教材之教育應用潛力_.md",
        "6.3 未來擴充建議.md",
        "6.4 發表與推廣建議.md",
        "6.5 給未來平台設計者的建議.md",
    ]),
    ("參考文獻", [
        "參考文獻.md",
    ]),
    ("附錄 A 量表完整題項與信度", [
        "附錄 A 量表完整題項與信度.md",
    ]),
]

# ── 正則式 ───────────────────────────────────────────────────────────────
RE_HEADING    = re.compile(r'^(#{1,4})\s+(.*)')
RE_BLOCKQUOTE = re.compile(r'^>\s?(.*)$')
RE_BOLD_ONLY  = re.compile(r'^\*\*(.+?)\*\*\s*$')
RE_FIG_CAP    = re.compile(r'^圖\s*\d+[-−–]\d+')
RE_FIG_DESC   = re.compile(r'^[（(]圖\s*\d+')
RE_SUBSEC     = re.compile(r'^\d+\.\d+\.\d+')   # X.X.X
RE_SEC        = re.compile(r'^\d+\.\d+\s')       # X.X<空格>
RE_CHAPTER_NO = re.compile(r'^第[一二三四五六七八九十百]+章')
RE_CN_SUB     = re.compile(r'^[一二三四五六七八九十]+[、．]')
RE_CN_PAREN   = re.compile(r'^[（(][一二三四五六七八九十]+[）)]')   # （一）/（二）...
RE_NUM_SUB    = re.compile(r'^\d+[、．]')                              # 1、/2、...
RE_LIST       = re.compile(r'^[-*+]\s+(.*)')
RE_TABLE_ROW  = re.compile(r'^\|')
RE_TABLE_SEP  = re.compile(r'^\|[\s\-:]+\|')
RE_HR         = re.compile(r'^-{3,}$|^\*{3,}$')
RE_MD_IMAGE   = re.compile(r'^!\[(.*?)\]\((.*?)\)\s*$')   # ![caption](path)
RE_FIG_KEY    = re.compile(r'圖\s*(\d+)[-−–](\d+)([a-z]?)')   # 圖 3-1 / 圖 3-11a
RE_TABLE_CAP  = re.compile(r'^表\s*\d+[-−–]\d+')   # 表說起始：表 5-5
RE_TABLE_KEY  = re.compile(r'表\s*(\d+)[-−–](\d+)([a-z]?)')   # 表 5-5 / 表 5-5a

# ── 圖表編號管理 ─────────────────────────────────────────────────────────

class FigureRegistry:
    """蒐集所有圖標籤並依出現順序給予全文連續編號。"""
    def __init__(self):
        self.mapping = {}   # (chapter, num, letter) → 連續號

    def register(self, caption_text):
        m = RE_FIG_KEY.search(caption_text)
        if not m:
            return
        key = (m.group(1), m.group(2), m.group(3))
        if key not in self.mapping:
            self.mapping[key] = len(self.mapping) + 1

    def convert(self, text):
        """將文字中所有「圖 X-Y」替換為「圖 N」（依登記表）。"""
        if not text:
            return text
        def repl(m):
            key = (m.group(1), m.group(2), m.group(3))
            if key in self.mapping:
                return f'圖 {self.mapping[key]}'
            return m.group(0)
        return RE_FIG_KEY.sub(repl, text)


class TableRegistry:
    """每個 tablecaption 出現一次就給一個唯一序號（避免不同檔重用 表 X-Y 而被合併）。
    文字內參照（如「見表 5-6」）採用 first-seen mapping，沿用第一次出現的序號。
    """
    def __init__(self):
        self.sequence = []          # 依出現順序：每筆 (key, number)
        self.first_seen = {}        # key → 首次出現序號（供文字參照轉換）
        self._render_cursor = 0

    def register(self, caption_text):
        m = RE_TABLE_KEY.search(caption_text)
        if not m:
            return
        key = (m.group(1), m.group(2), m.group(3))
        n = len(self.sequence) + 1
        self.sequence.append((key, n))
        self.first_seen.setdefault(key, n)

    def render_caption_number(self, caption_text):
        """渲染期 — 為當前 tablecaption 取出對應序號（按事件順序消耗 sequence）。"""
        m = RE_TABLE_KEY.search(caption_text)
        if not m:
            return caption_text
        if self._render_cursor >= len(self.sequence):
            return caption_text
        _, n = self.sequence[self._render_cursor]
        self._render_cursor += 1
        return RE_TABLE_KEY.sub(lambda _: f'表 {n}', caption_text, count=1)

    def convert(self, text):
        """文字參照 → first-seen 序號。"""
        if not text:
            return text
        def repl(m):
            key = (m.group(1), m.group(2), m.group(3))
            if key in self.first_seen:
                return f'表 {self.first_seen[key]}'
            return m.group(0)
        return RE_TABLE_KEY.sub(repl, text)

    @property
    def mapping(self):
        """供 build() 摘要報告使用。"""
        return {k: v for k, v in self.sequence}


def normalize_caption_punct(text):
    """統一圖/表 caption 的編號與標題之間用全形空格分隔。
    「圖 1：xxx」「圖 1 xxx」「表 5  xxx」 → 「圖 1　xxx」「表 5　xxx」
    """
    if not text:
        return text
    # 「圖 N」或「表 N」 後若接 半形空格/全形空格/半形冒號/全形冒號 + 標題文字 → 替換成單一全形空格
    text = re.sub(r'(圖\s*\d+[a-z]?)[\s::　]+(?=\S)', r'\1　', text)
    text = re.sub(r'(表\s*\d+[a-z]?)[\s::　]+(?=\S)', r'\1　', text)
    return text

# ── 字型 / 段落格式函數 ──────────────────────────────────────────────────

def set_font(run, size_pt, bold=False):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = FONT_EN
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    rFonts.set(qn('w:ascii'),    FONT_EN)
    rFonts.set(qn('w:hAnsi'),   FONT_EN)

def set_spacing(para, line_mult=1.2, before_pt=0, after_pt=0):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'),     str(int(240 * line_mult)))
    sp.set(qn('w:lineRule'), 'auto')
    if before_pt:
        sp.set(qn('w:before'), str(int(before_pt * 20)))
    if after_pt:
        sp.set(qn('w:after'), str(int(after_pt * 20)))
    pPr.append(sp)

def set_first_line_indent_chars(para, chars):
    """中文段落首行縮排 N 字元（firstLineChars 以百分位寫入）。"""
    if chars <= 0:
        return
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:ind')):
        pPr.remove(old)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLineChars'), str(chars * 100))
    # fallback: 14pt × N 字 × 20 twips/pt
    ind.set(qn('w:firstLine'), str(int(14 * chars * 20)))
    pPr.append(ind)

def _add_field(run, instr_text):
    """在 run 內加入一個 Word field（如 PAGE / TOC）。"""
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instr_text
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)

def add_page_number_field(footer_para):
    """在 footer 段落加入 PAGE field（顯示當前頁碼）。"""
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    set_font(run, 12)
    _add_field(run, 'PAGE')

def set_section_page_format(section, fmt='decimal', start=1):
    """設定 section 的頁碼格式：'decimal'（1,2,3）或 'upperRoman'（I,II,III）。"""
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    pg = OxmlElement('w:pgNumType')
    pg.set(qn('w:fmt'), fmt)
    pg.set(qn('w:start'), str(start))
    sectPr.append(pg)

def add_section_break(doc, page_fmt='decimal', start=1):
    """在文件尾插入 section break (next page) 並設定頁碼格式。"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    # 沿用上一段邊距
    last_sect = doc.sections[-1]._sectPr
    for tag in ('w:pgSz', 'w:pgMar', 'w:cols', 'w:docGrid'):
        elem = last_sect.find(qn(tag))
        if elem is not None:
            from copy import deepcopy
            sectPr.append(deepcopy(elem))
    pg_type = OxmlElement('w:type')
    pg_type.set(qn('w:val'), 'nextPage')
    sectPr.insert(0, pg_type)
    pg = OxmlElement('w:pgNumType')
    pg.set(qn('w:fmt'), page_fmt)
    pg.set(qn('w:start'), str(start))
    sectPr.append(pg)
    pPr.append(sectPr)

RE_INLINE_MATH = re.compile(r'\$([^$\n]+?)\$')

_STAR_TOKEN = '★'  # 罕見 unicode，用來暫代 markdown 內的 \* escape

def _restore_md_escapes(s):
    """把 add_runs 開頭塞進去的 token 還原成字面 *。"""
    return s.replace(_STAR_TOKEN, '*')

def add_runs(para, text, size_pt, default_bold=False):
    """拆解 **bold** 與 $inline_math$ 標記，加入 runs / OMML。
    處理 markdown escape：`\\*` 表示字面 *，需先以 token 暫代避免被 ** bold 邏輯誤切。
    """
    # 預處理：md escape `\*` → token，等切完 bold/math 後再還原成 *
    text = text.replace(r'\*', _STAR_TOKEN)
    # 先按 bold 切，再對每個非 bold 段切 inline math
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        is_bold = default_bold or (i % 2 == 1)
        if is_bold:
            # bold 段內也可能有 math，但暫不處理（罕見）
            run = para.add_run(_restore_md_escapes(part))
            set_font(run, size_pt, bold=True)
            continue
        # 切 inline math
        last_end = 0
        for m in RE_INLINE_MATH.finditer(part):
            if m.start() > last_end:
                run = para.add_run(_restore_md_escapes(part[last_end:m.start()]))
                set_font(run, size_pt, bold=False)
            omml = latex_to_omml(m.group(1), display=False)
            if omml is not None:
                para._p.append(omml)
            else:
                run = para.add_run(_restore_md_escapes(m.group(0)))
                set_font(run, size_pt, bold=False)
            last_end = m.end()
        if last_end < len(part):
            run = para.add_run(_restore_md_escapes(part[last_end:]))
            set_font(run, size_pt, bold=False)

# ── 段落類型函數 ─────────────────────────────────────────────────────────

def _apply_heading_style(para, level):
    """套用內建 Heading 樣式（讓 TOC field 能抓到）。"""
    try:
        para.style = f'Heading {level}'
    except KeyError:
        pass

def add_chapter_title(doc, title):
    """章名：標楷體 20pt 置中 1.2x 前後留雙倍行距。"""
    p = doc.add_paragraph()
    _apply_heading_style(p, 1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.2, before_pt=28, after_pt=28)
    run = p.add_run(title)
    set_font(run, 20, bold=True)

def add_section(doc, text):
    """X.X 節名：16pt 靠左。"""
    p = doc.add_paragraph()
    _apply_heading_style(p, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, 1.2, before_pt=14, after_pt=0)
    run = p.add_run(text)
    set_font(run, 16, bold=True)

def add_subsection(doc, text):
    """X.X.X 節名：14pt 靠左。"""
    p = doc.add_paragraph()
    _apply_heading_style(p, 3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, 1.2, before_pt=10, after_pt=0)
    run = p.add_run(text)
    set_font(run, 14, bold=True)

def add_subhead(doc, text):
    """一、 層次：14pt 靠左 粗體。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, 1.2)
    run = p.add_run(text)
    set_font(run, 14, bold=True)

def add_body(doc, text):
    """正文：14pt justified 1.2x，首行縮排 2 字元。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, 1.2)
    set_first_line_indent_chars(p, BODY_FIRST_LINE_INDENT_CHARS)
    add_runs(p, text, 14)

def add_body_15x(doc, text):
    """1.5 倍行距正文（摘要 / 誌謝 / 參考文獻使用）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, 1.5)
    set_first_line_indent_chars(p, BODY_FIRST_LINE_INDENT_CHARS)
    add_runs(p, text, 14)

def add_toc_field(doc, title, switches):
    """加入 TOC field（Word 開檔時按 F9 即可生成）。
    switches 範例：'\\o "1-3" \\h \\z \\u'（目錄）、'\\c "圖"'（圖目錄）、'\\c "表"'（表目錄）
    """
    add_chapter_title(doc, title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, 1.5)
    run = p.add_run()
    set_font(run, 14)
    _add_field(run, f'TOC {switches}')
    # 提示文字（field 渲染前會先顯示）
    hint = doc.add_paragraph()
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(hint, 1.2)
    hr = hint.add_run('（請於 Word 中按 F9 更新目錄）')
    set_font(hr, 10)

def add_quote(doc, text):
    """引言段落（markdown `> ...`）：左縮排、無首行縮排，視覺上區分於正文。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, 1.2, before_pt=2, after_pt=2)
    # 左縮排 2 字，去掉首行縮排
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:ind')):
        pPr.remove(old)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:leftChars'), '200')   # 整段左縮 2 字
    ind.set(qn('w:left'), str(2 * 14 * 20))
    ind.set(qn('w:firstLineChars'), '0')
    ind.set(qn('w:firstLine'), '0')
    pPr.append(ind)
    add_runs(p, text, 13)

def add_list_item(doc, text, line_mult=1.2):
    """清單項目：• 開頭，14pt；可指定行距（參考文獻章用 1.5x）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, line_mult)
    run = p.add_run('• ')
    set_font(run, 14)
    add_runs(p, text, 14)

def add_fig_placeholder(doc, caption_text):
    """圖片不存在時的佔位文字。"""
    p = doc.add_paragraph(f'【圖片待補：{caption_text}】')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.2, before_pt=8, after_pt=4)
    for run in p.runs:
        set_font(run, 12)
        run.font.color.rgb = None  # 用預設色

def add_fig_caption(doc, text):
    """圖說：12pt 置中，置於圖下方。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.2, before_pt=4, after_pt=8)
    run = p.add_run(text)
    set_font(run, 12, bold=True)

def add_table_caption(doc, text):
    """表說：12pt 粗體置中，置於表上方（規範要求表號表名列於表上方）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.2, before_pt=8, after_pt=4)
    run = p.add_run(text)
    set_font(run, 12, bold=True)

def add_block_math(doc, tex):
    """LaTeX block 方程 → 獨立段落，OMML 置中顯示。
    display=True 時 XSL 會直接輸出 <m:oMathPara><m:oMath>...</m:oMath></m:oMathPara>，
    直接 append 即可，不可再額外包一層 oMathPara（會造成 Word 拒絕渲染）。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.2, before_pt=4, after_pt=4)
    omml = latex_to_omml(tex, display=True)
    if omml is None:
        run = p.add_run(f'$${tex}$$')
        set_font(run, 14)
        return
    p._p.append(omml)

def add_image(doc, img_path, width_cm=14):
    """插入圖片，置中。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.0, before_pt=8)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def add_table(doc, raw_rows, fig_registry=None, table_registry=None):
    """從 md 表格列插入 Word 表格；registry 提供時，cell 內圖/表參照會轉成新編號。

    欄數越多，字級越小，避免在 A4 寬度下單欄被擠到字斷行：
      ≤ 6 欄 → 12pt（標準正文字級）
      7 欄   → 11pt
      8 欄   → 10pt
      ≥ 9 欄 → 9pt
    """
    rows = [r for r in raw_rows if not RE_TABLE_SEP.match(r.strip())]
    if not rows:
        return
    cols = max(len(r.strip('|').split('|')) for r in rows)
    if cols <= 0:
        return
    # 依欄數決定字級
    if cols <= 6:
        font_size = 12
    elif cols == 7:
        font_size = 11
    elif cols == 8:
        font_size = 10
    else:
        font_size = 9
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = 'Table Grid'
    tbl.autofit = True
    for ri, row_line in enumerate(rows):
        cells = [c.strip() for c in row_line.strip('|').split('|')]
        for ci in range(cols):
            cell_text = cells[ci].strip('*').strip() if ci < len(cells) else ''
            if fig_registry is not None:
                cell_text = fig_registry.convert(cell_text)
            if table_registry is not None:
                cell_text = table_registry.convert(cell_text)
            cell = tbl.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_spacing(p, 1.2)
            add_runs(p, cell_text, font_size)

# ── 圖片查找 ─────────────────────────────────────────────────────────────

def find_image(caption):
    """從圖說文字找對應圖片，例：圖 3-1 → images/fig3-1.png，圖 3-11a → images/fig3-11a.png。"""
    m = re.search(r'圖\s*(\d+)\s*[-−–]\s*(\d+)([a-z]?)', caption)
    if not m:
        return None
    ch, num, letter = m.group(1), m.group(2), m.group(3)
    for ext in ('png', 'jpg', 'jpeg'):
        p = IMG_DIR / f"fig{ch}-{num}{letter}.{ext}"
        if p.exists():
            return p
    return None

def resolve_md_image(path_str):
    """解析 markdown 圖片相對路徑。md 在 MD_DIR/，所以 ../images/x.png → IMG_DIR/x.png。"""
    p = (MD_DIR / path_str).resolve()
    if p.exists():
        return p
    return None

# ── MD 解析 ──────────────────────────────────────────────────────────────

def normalize(text):
    """去除 ** 標記並 strip。"""
    return text.replace('**', '').strip()

def classify_bold_line(content):
    """判斷純 bold 行的段落類型（圖說 / 表說 / 中文層次）。
    md 規範已要求節 / 子節用 ##/### heading；此處只處理圖表標題與「一、」層次。
    """
    if RE_FIG_CAP.match(content):
        return 'figcaption'
    if RE_TABLE_CAP.match(content):
        return 'tablecaption'
    if RE_CN_SUB.match(content) or RE_CN_PAREN.match(content) or RE_NUM_SUB.match(content):
        return 'subhead'
    return 'body'

def parse_md(filepath, chapter_title):
    """逐行解析 md，yield (kind, content)。

    md 規範（docs/md_writing_spec.md）要求：
      - 節用 `## X.X 標題`、子節用 `### X.X.X 標題`
      - 圖片用 `![圖 X-Y　標題](路徑)`，標題不重複貼一次 bold
      - 表格上方用 `**表 X-Y　標題**`
    """
    lines = filepath.read_text(encoding='utf-8').splitlines()
    i = 0
    last_md_img_key = None  # 上一個 md_image 的圖編號 (ch, num, letter)；遇非空非 figcaption 內容即清除
    while i < len(lines):
        line = lines[i]
        i += 1
        stripped = line.strip()

        if not stripped or RE_HR.match(stripped):
            continue  # 空行/分隔線：保留 last_md_img_key 以容許段落間空行

        # LaTeX block math: $$ ... $$
        if stripped == '$$':
            tex_lines = []
            while i < len(lines) and lines[i].strip() != '$$':
                tex_lines.append(lines[i])
                i += 1
            i += 1  # skip closing $$
            last_md_img_key = None
            yield ('block_math', '\n'.join(tex_lines).strip())
            continue

        # 舊式圖描述 （圖 X-X ...） → 跳過
        if RE_FIG_DESC.match(stripped):
            continue

        # Markdown 圖片語法 ![caption](path)
        m = RE_MD_IMAGE.match(stripped)
        if m:
            caption = m.group(1).strip()
            mk = RE_FIG_KEY.search(caption)
            last_md_img_key = (mk.group(1), mk.group(2), mk.group(3)) if mk else None
            yield ('md_image', (caption, m.group(2).strip()))
            continue

        # 表格
        if RE_TABLE_ROW.match(stripped):
            rows = [line]
            while i < len(lines) and RE_TABLE_ROW.match(lines[i].strip()):
                rows.append(lines[i])
                i += 1
            last_md_img_key = None
            yield ('table', rows)
            continue

        # Markdown heading
        m = RE_HEADING.match(line)
        if m:
            content = normalize(m.group(2))
            if not content:
                continue
            if content.replace('　', ' ') == chapter_title.replace('　', ' '):
                continue
            last_md_img_key = None
            if RE_SUBSEC.match(content):
                yield ('subsection', content)
            elif RE_SEC.match(content):
                yield ('section', content)
            elif RE_CN_SUB.match(content) or RE_CN_PAREN.match(content) or RE_NUM_SUB.match(content):
                yield ('subhead', content)
            else:
                yield ('section', content)
            continue

        # 純 **bold** 行（圖說 / 表說 / 一、層次）
        m = RE_BOLD_ONLY.match(stripped)
        if m:
            content = m.group(1).strip()
            kind = classify_bold_line(content)
            # 去重：md_image 後緊接同編號 figcaption → 跳過
            if kind == 'figcaption' and last_md_img_key is not None:
                mk = RE_FIG_KEY.search(content)
                if mk and (mk.group(1), mk.group(2), mk.group(3)) == last_md_img_key:
                    last_md_img_key = None
                    continue
            last_md_img_key = None
            yield (kind, content)
            continue

        # 清單
        m = RE_LIST.match(stripped)
        if m:
            last_md_img_key = None
            yield ('list', m.group(1))
            continue

        # Markdown blockquote `> ...` → 剝掉 `>` 當引言段落輸出
        m = RE_BLOCKQUOTE.match(stripped)
        if m:
            last_md_img_key = None
            yield ('quote', m.group(1))
            continue

        # 正文
        last_md_img_key = None
        yield ('body', stripped)

# ── 主程式 ───────────────────────────────────────────────────────────────

def collect_events():
    """Pass 1：掃描所有 md，收集 (chapter_title, kind, content) 序列。"""
    events = []
    for chapter_title, files in CHAPTERS:
        events.append((chapter_title, '__chapter_start__', None))
        for fname in files:
            fpath = MD_DIR / fname
            if not fpath.exists():
                events.append((chapter_title, '__missing_file__', fname))
                continue
            events.append((chapter_title, '__file_start__', fname))
            for kind, content in parse_md(fpath, chapter_title):
                events.append((chapter_title, kind, content))
    return events

def build_fig_registry(events):
    """Pass 2：依出現順序登記所有圖標籤。"""
    reg = FigureRegistry()
    for _, kind, content in events:
        if kind == 'figcaption':
            reg.register(content)
        elif kind == 'md_image':
            caption, _ = content
            reg.register(caption)
    return reg

def build_table_registry(events):
    """Pass 2b：依出現順序登記所有表標籤。"""
    reg = TableRegistry()
    for _, kind, content in events:
        if kind == 'tablecaption':
            reg.register(content)
    return reg

def _add_centered_line(doc, text, size_pt=14, bold=False, before_pt=0, after_pt=0, line_mult=1.5):
    """前置頁專用：置中段落，可控字級 / 粗體 / 段距。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, line_mult, before_pt=before_pt, after_pt=after_pt)
    run = p.add_run(text)
    set_font(run, size_pt, bold=bold)
    return p

def add_cover_page(doc):
    """元智格式 A 附件一 封面（R-G-A1）。"""
    M = THESIS_META
    # 上方留白
    _add_centered_line(doc, '', 14, before_pt=60)
    _add_centered_line(doc, '元　智　大　學', 24, bold=True, after_pt=24)
    _add_centered_line(doc, M['dept_zh'], 22, bold=True, after_pt=24)
    _add_centered_line(doc, f"{M['degree_zh']}　論　文", 22, bold=True, after_pt=60)
    _add_centered_line(doc, M['title_zh'], 20, bold=True, after_pt=80)
    _add_centered_line(doc, f"研  究  生：{M['student_zh']}", 16, after_pt=12)
    _add_centered_line(doc, f"指 導 教 授：{M['advisor_zh']}　博士", 16, after_pt=80)
    _add_centered_line(doc, f"中 華 民 國 　{M['year_roc']}　年　{M['month_zh']}　月", 16)
    add_page_break(doc)

def add_title_page(doc):
    """元智格式 A 附件二 書名頁（R-G-A2）。整頁須一頁容納。"""
    M = THESIS_META
    _add_centered_line(doc, M['title_zh'], 18, bold=True, before_pt=24, after_pt=6, line_mult=1.2)
    _add_centered_line(doc, M['title_en'].upper(), 14, bold=True, after_pt=20, line_mult=1.2)

    # 研究生 / 指導教授（中英對照）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.2, after_pt=4)
    r1 = p.add_run(f"研  究  生：{M['student_zh']}")
    set_font(r1, 14)
    r2 = p.add_run(f"      Student：{M['student_en']}")
    set_font(r2, 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 1.2, after_pt=20)
    r1 = p.add_run(f"指 導 教 授：{M['advisor_zh']}")
    set_font(r1, 14)
    r2 = p.add_run(f"      Advisor：{M['advisor_en']}")
    set_font(r2, 14)

    _add_centered_line(doc, '元  智  大  學', 16, bold=True, after_pt=4, line_mult=1.2)
    _add_centered_line(doc, M['dept_zh'], 16, bold=True, after_pt=4, line_mult=1.2)
    _add_centered_line(doc, f"{M['degree_zh']}  論  文", 16, bold=True, after_pt=18, line_mult=1.2)

    _add_centered_line(doc, 'A Thesis', 14, after_pt=2, line_mult=1.2)
    _add_centered_line(doc, f"Submitted to Department of {M['dept_en']}", 14, after_pt=2, line_mult=1.2)
    _add_centered_line(doc, M['college_en'], 14, after_pt=2, line_mult=1.2)
    _add_centered_line(doc, 'Yuan Ze University', 14, after_pt=2, line_mult=1.2)
    _add_centered_line(doc, 'in Partial Fulfillment of the Requirements', 14, after_pt=2, line_mult=1.2)
    _add_centered_line(doc, 'for the Degree of', 14, after_pt=2, line_mult=1.2)
    _add_centered_line(doc, M['degree_en'], 14, after_pt=2, line_mult=1.2)
    _add_centered_line(doc, 'in', 14, after_pt=2, line_mult=1.2)
    _add_centered_line(doc, M['dept_en'], 14, after_pt=10, line_mult=1.2)

    _add_centered_line(doc, f"{M['month_en']} {M['year_ad']}", 14, after_pt=4, line_mult=1.2)
    _add_centered_line(doc, M['location_en'], 14, after_pt=10, line_mult=1.2)
    _add_centered_line(doc, f"中華民國　{M['year_roc']}　年　{M['month_zh']}　月", 14, line_mult=1.2)
    add_page_break(doc)

def add_verification_page(doc):
    """元智格式 A 附件三 論文口試委員審定書（R-G-A3）。
    口試後委員簽名版才能附正本；此處先放正式格式的空白表，方便列印簽署。
    """
    M = THESIS_META
    _add_centered_line(doc, '元智大學碩士班研究生', 18, bold=True, after_pt=6)
    _add_centered_line(doc, '論文口試委員審定書', 18, bold=True, after_pt=12)
    _add_centered_line(doc, 'YZU Master Program', 14, after_pt=2)
    _add_centered_line(doc, 'Verification Letter from the Oral Examination Committee', 14, after_pt=30)

    # 學年期
    _add_centered_line(doc, '第　一一四　學年度第　二　學期', 14, after_pt=24)
    _add_centered_line(doc, f"{M['dept_zh']}　{M['student_zh']}　君，學號 {M['student_id']}", 14, after_pt=12)
    _add_centered_line(doc, '所提之論文', 14, after_pt=12)
    _add_centered_line(doc, M['title_zh'], 14, bold=True, after_pt=24)
    _add_centered_line(doc, f"經本委員會審議，認為符合 {M['degree_zh']} 資格標準。", 14, after_pt=36)

    _add_centered_line(doc, f"Department of {M['dept_en']}", 14, after_pt=4)
    _add_centered_line(doc, f"Name: {M['student_en']}     Student No.: {M['student_id']}", 14, after_pt=4)
    _add_centered_line(doc, f"Thesis: {M['title_en']}", 14, after_pt=12)
    _add_centered_line(doc, "Fulfill the Master's Degree's Requirement; reviewed and verified by this committee.", 14, after_pt=36)

    # 簽名欄（靠左）
    for label in ['口試委員 Oral Examination Committee Members：', '指 導 教 授 Advisor：', '系  主  任 Head of Dept.：']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, 2.0, after_pt=20)
        run = p.add_run(label)
        set_font(run, 14)

    _add_centered_line(doc, f"中華民國　{M['year_roc']}　年(Year)　{M['month_zh']}　月(Month)　　　日(Date)", 14, before_pt=20)
    add_page_break(doc)


def _collect_toc_entries(events, fig_registry, table_registry):
    """掃 events 收集三組目錄條目（不含頁碼）。
    每條目附 sequence number 區分（如多章節都有「伍、小結」）。
    """
    chap_entries = []     # [(level, text, seq)]
    fig_entries  = []     # [(num, caption_text)]
    table_entries = []
    chap_seq = 0
    fig_counter = 0
    tab_counter = 0

    for chapter_title, kind, content in events:
        if kind == '__chapter_start__':
            chap_seq += 1
            chap_entries.append((1, chapter_title, chap_seq))
        elif kind == 'section':
            text = convert_section_text(content) if CONVERT_SECTION_NUMBERING else content
            chap_seq += 1
            chap_entries.append((2, text, chap_seq))
        elif kind == 'subsection':
            text = convert_subsection_text(content) if CONVERT_SUBSECTION_NUMBERING else content
            chap_seq += 1
            chap_entries.append((3, text, chap_seq))
        elif kind == 'figcaption':
            fig_counter += 1
            # 用 fig_registry 把原號轉成全文編號顯示
            display = fig_registry.convert(content) if fig_registry else content
            display = normalize_caption_punct(display) if NORMALIZE_CAPTION_PUNCT else display
            fig_entries.append((fig_counter, display))
        elif kind == 'md_image':
            caption, _ = content
            fig_counter += 1
            display = fig_registry.convert(caption) if fig_registry else caption
            display = normalize_caption_punct(display) if NORMALIZE_CAPTION_PUNCT else display
            fig_entries.append((fig_counter, display))
        elif kind == 'tablecaption':
            tab_counter += 1
            # tablecaption 渲染期用 render_caption_number 取序號；這裡用相同 sequence
            if table_registry and table_registry.sequence:
                idx = min(tab_counter - 1, len(table_registry.sequence) - 1)
                _, n = table_registry.sequence[idx]
                display = RE_TABLE_KEY.sub(lambda _m: f'表 {n}', content, count=1)
            else:
                display = content
            display = normalize_caption_punct(display) if NORMALIZE_CAPTION_PUNCT else display
            table_entries.append((tab_counter, display))

    return chap_entries, fig_entries, table_entries


def add_static_toc(doc, title, entries, kind='chap'):
    """寫一份靜態目錄（不依賴 Word TOC field）。
    entries: 章節為 [(level, text)]；圖/表為 [(num, caption)]。
    kind: 'chap' / 'fig' / 'table'，控制條目樣式。
    每條目右側以 dot-leader tab 對齊頁碼（頁碼由 toc_pages.json 查得，
    找不到顯示「—」，待 2-pass 工作流填入）。
    """
    add_chapter_title(doc, title)
    if not entries:
        return

    PAGE_TAB_CM = 15.5  # 條目右側頁碼欄對齊位置（左邊距內側 ~15.5cm）

    for item in entries:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, 1.5, before_pt=0, after_pt=0)

        # 加 dot-leader right tab stop
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(PAGE_TAB_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )

        if kind == 'chap':
            level, text, seq = item
            indent = (level - 1) * 2   # 章=0 / 節=2 / 小節=4 字
            if indent:
                pPr = p._p.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:leftChars'), str(indent * 100))
                ind.set(qn('w:left'), str(indent * 14 * 20))
                pPr.append(ind)
            size = 14 if level == 1 else 13
            run = p.add_run(text)
            set_font(run, size, bold=(level == 1))
            page_str = _toc_page(f'CHAP::{seq}::{text}')
            run_p = p.add_run('\t' + page_str)
            set_font(run_p, size, bold=(level == 1))
        else:
            num, caption = item
            run = p.add_run(caption)
            set_font(run, 12)
            prefix = 'FIG' if kind == 'fig' else 'TABLE'
            page_str = _toc_page(f'{prefix}::{caption[:40]}')
            run_p = p.add_run('\t' + page_str)
            set_font(run_p, 12)


def add_front_matter(doc, events=None, fig_registry=None, table_registry=None):
    """加入前置部分（封面 / 書名頁 / 審定書 / 摘要 / 誌謝 / 目錄 / 圖目錄 / 表目錄）。
    封面/書名頁/審定書依元智格式 A 附件一/二/三填入；摘要與誌謝讀 md/。
    目錄三組（目錄 / 表目錄 / 圖目錄）為靜態列表，不依賴 Word TOC field。
    """
    add_cover_page(doc)
    add_title_page(doc)
    add_verification_page(doc)

    # 中文摘要 + 英文 Abstract（兩個 ## section，第二個前加換頁）
    md_path = MD_DIR / '摘要.md'
    if md_path.exists():
        section_count = 0
        # 傳空字串作為 chapter_title，使 ## 中文摘要 / ## Abstract 兩個 heading 都被 yield
        for kind, content in parse_md(md_path, ''):
            if kind == 'section':
                section_count += 1
                if section_count > 1:
                    add_page_break(doc)
                # 摘要標題置中、字級略大、上下留白
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_spacing(p, 1.5, before_pt=0, after_pt=18)
                run = p.add_run(content)
                set_font(run, 18, bold=True)
            elif kind == 'body':
                add_body_15x(doc, content)
    add_page_break(doc)

    # 誌謝
    md_path = MD_DIR / '誌謝.md'
    if md_path.exists():
        for kind, content in parse_md(md_path, '誌謝'):
            if kind == 'section':
                add_section(doc, content)
            elif kind == 'body':
                add_body_15x(doc, content)
    add_page_break(doc)

    # 目錄 / 表目錄 / 圖目錄（依規範 R-G-A9_目錄中 順序）
    if events is not None:
        chap_e, fig_e, table_e = _collect_toc_entries(events, fig_registry, table_registry)
    else:
        chap_e, fig_e, table_e = [], [], []
    add_static_toc(doc, '目錄',   chap_e,  kind='chap')
    add_page_break(doc)
    add_static_toc(doc, '表目錄', table_e, kind='table')
    add_page_break(doc)
    add_static_toc(doc, '圖目錄', fig_e,   kind='fig')


def build():
    doc = Document()

    # 頁面設定（A4 + 邊距 + footer 距離）
    for sec in doc.sections:
        sec.page_width      = Cm(21.0)   # A4 寬
        sec.page_height     = Cm(29.7)   # A4 高
        sec.top_margin      = Cm(3.5)
        sec.bottom_margin   = Cm(2.0)
        sec.left_margin     = Cm(4.0)
        sec.right_margin    = Cm(2.0)
        sec.footer_distance = Cm(1.0)    # 規範：版面底端 1cm 處中央繕打頁次

    # footer 中央頁碼
    if ADD_PAGE_NUMBER:
        for sec in doc.sections:
            footer_para = sec.footer.paragraphs[0]
            footer_para.text = ''
            add_page_number_field(footer_para)

    # 清除預設段落間距
    normal = doc.styles['Normal']
    normal.paragraph_format.space_after  = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    # 強制 Word 開檔時自動更新所有 field（含 TOC / PAGE）
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')

    # Pass 1: 先收集事件（給靜態目錄使用）
    events = collect_events()

    # Pass 2: 建立圖/表編號對照
    fig_registry   = build_fig_registry(events)   if CONVERT_FIGURE_NUMBERING else None
    table_registry = build_table_registry(events) if CONVERT_TABLE_NUMBERING  else None

    # 前置部分（小寫羅馬數字 i, ii, iii...）— 靜態目錄需要 events + registries
    # OOXML 行為：inline sectPr（add_section_break 插入的）控制【之前】的內容（前置）；
    # 文件尾的 sectPr（doc.sections[-1]）控制【之後】（即正文）。
    # 注意 python-docx 加入 inline sectPr 後，doc.sections 會有兩個 section；
    # [0] 是 inline（前置）、[-1] 是 trailing（正文），必須用 [-1] 設正文格式。
    add_front_matter(doc, events=events, fig_registry=fig_registry, table_registry=table_registry)
    add_section_break(doc, page_fmt='lowerRoman', start=1)              # 前置 → 小寫羅馬
    set_section_page_format(doc.sections[-1], fmt='decimal', start=1)   # 正文 → 阿拉伯數字

    # Pass 3: 渲染 docx
    first_chapter = True
    missing_images = []
    use_15x_spacing = False  # 參考文獻章用 1.5x

    def maybe_convert_refs(text):
        """文內若有圖/表參照，依登記表替換；並統一 caption 標點。"""
        if not text:
            return text
        if fig_registry is not None:
            text = fig_registry.convert(text)
        if table_registry is not None:
            text = table_registry.convert(text)
        if NORMALIZE_CAPTION_PUNCT:
            text = normalize_caption_punct(text)
        return text

    for chapter_title, kind, content in events:
        if kind == '__chapter_start__':
            if not first_chapter:
                add_page_break(doc)
            first_chapter = False
            use_15x_spacing = (chapter_title == '參考文獻')
            print(f'\n▶ {chapter_title}')
            add_chapter_title(doc, chapter_title)
        elif kind == '__missing_file__':
            print(f'  ⚠  找不到：{content}')
        elif kind == '__file_start__':
            print(f'  ✓  {content}')
        elif kind == 'chapter_skip':
            pass
        elif kind == 'section':
            text = convert_section_text(content) if CONVERT_SECTION_NUMBERING else content
            add_section(doc, text)
        elif kind == 'subsection':
            text = convert_subsection_text(content) if CONVERT_SUBSECTION_NUMBERING else content
            add_subsection(doc, text)
        elif kind == 'subhead':
            add_subhead(doc, content)
        elif kind == 'figcaption':
            display = maybe_convert_refs(content)
            img = find_image(content)
            if img:
                add_image(doc, img)
            else:
                add_fig_placeholder(doc, display)
                missing_images.append(content)
            add_fig_caption(doc, display)
        elif kind == 'md_image':
            caption, path = content
            display = maybe_convert_refs(caption)
            img = resolve_md_image(path) or find_image(caption)
            if img:
                add_image(doc, img)
            else:
                add_fig_placeholder(doc, display)
                missing_images.append(caption)
            add_fig_caption(doc, display)
        elif kind == 'tablecaption':
            # caption 自己用渲染序號（避免 first-seen 把後續同名表壓回第一張）
            if table_registry is not None:
                display = table_registry.render_caption_number(content)
            else:
                display = content
            if NORMALIZE_CAPTION_PUNCT:
                display = normalize_caption_punct(display)
            add_table_caption(doc, display)
        elif kind == 'list':
            add_list_item(doc, maybe_convert_refs(content), line_mult=1.5 if use_15x_spacing else 1.2)
        elif kind == 'quote':
            add_quote(doc, maybe_convert_refs(content))
        elif kind == 'table':
            add_table(doc, content, fig_registry, table_registry)
        elif kind == 'block_math':
            add_block_math(doc, content)
        elif kind == 'body':
            if content:
                if use_15x_spacing:
                    add_body_15x(doc, maybe_convert_refs(content))
                else:
                    add_body(doc, maybe_convert_refs(content))

    doc.save(str(OUTPUT))
    print(f'\n✅ 輸出：{OUTPUT}')
    print(f'   ※ 開檔時 Word 會跳出「是否更新欄位」對話框 → 點「是」即可填入目錄與頁碼')

    # 摘要報告
    if fig_registry is not None and fig_registry.mapping:
        print(f'\n📊 圖編號對照（共 {len(fig_registry.mapping)} 張）：')
        for (ch, num, letter), new_n in sorted(fig_registry.mapping.items(), key=lambda x: x[1]):
            old = f'圖 {ch}-{num}{letter}'
            print(f'   {old:>12s}  →  圖 {new_n}')

    if table_registry is not None and table_registry.mapping:
        print(f'\n📋 表編號對照（共 {len(table_registry.mapping)} 張）：')
        for (ch, num, letter), new_n in sorted(table_registry.mapping.items(), key=lambda x: x[1]):
            old = f'表 {ch}-{num}{letter}'
            print(f'   {old:>12s}  →  表 {new_n}')

    if missing_images:
        print(f'\n⚠  以下 {len(missing_images)} 張圖片未找到，已插入佔位文字：')
        for caption in missing_images:
            print(f'   - {caption}')
        print(f'\n   圖片請命名為 images/fig{{章}}-{{號}}.png')
        print(f'   例：圖 3-1 → images/fig3-1.png')

if __name__ == '__main__':
    build()
